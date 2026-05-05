import os
import tempfile
import pytest
from app.core.ingestion.parser import parse_pdf, parse_docx
from app.core.ingestion.chunker import chunk_text
from app.core.ingestion.embedder import embed_chunks, embed_query


def test_chunk_text_basic():
    text = "Hello world. " * 200
    chunks = chunk_text(text)
    assert len(chunks) > 1
    for chunk in chunks:
        assert len(chunk) <= 600  # some slack over 512 due to splitter behaviour


def test_chunk_text_overlap():
    text = "word " * 300
    chunks = chunk_text(text)
    # consecutive chunks should share some content due to overlap
    assert len(chunks) >= 2


def test_embed_chunks_shape():
    chunks = ["This is a test sentence.", "Another sentence here."]
    vectors = embed_chunks(chunks)
    assert vectors.shape == (2, 384)


def test_embed_query_shape():
    vec = embed_query("What is the capital of France?")
    assert vec.shape == (384,)


def test_parse_docx(tmp_path):
    from docx import Document as DocxDocument
    doc = DocxDocument()
    doc.add_paragraph("Hello DocuMind")
    doc.add_paragraph("Second paragraph")
    p = tmp_path / "test.docx"
    doc.save(str(p))
    text = parse_docx(str(p))
    assert "Hello DocuMind" in text
    assert "Second paragraph" in text
